import os
import sys
import collections

from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import QMessageBox
import pyqtgraph as pg
import PGQ_image_exporter # modified pyqtgraph Exporter!
import numpy as np
import ui
import form
from cmath import rect
from docx import Document
import datetime
# Команда для преобразования файла QtDesigner в питоновский:
# -------->  pyuic5 UI.ui -o ui.py
# -------->  pyuic5 form.ui -o form.py


class Chart(QtWidgets.QDialog, form.Ui_Dialog):

    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.openButton.clicked.connect(self.plot_to_word)

    def plot_to_word(self):

        # create an exporter instance, as an argument give it
        # the item you wish to export
        exporter = PGQ_image_exporter.PQG_ImageExporter(self.widget.plotItem)

        # set export parameters if needed
        exporter.parameters()['width'] = 600

        # save to file
        exporter.export('fileName.png')



        document = Document('test.docx')
        document.add_paragraph(
            'Lorem ipsum dolor sit amet. {}'.format(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        document.add_picture('fileName.png')
        document.add_page_break()
        document.save('test.docx')
        Chart.display_ok_message('Word saved!')

    @classmethod
    def display_ok_message(cls, message):
        msg = QtWidgets.QMessageBox()
        msg.setIcon(QMessageBox.Information)

        msg.setText(message)
        msg.setInformativeText("It's ok!")
        msg.setWindowTitle("MessageBox title")
        # msg.setDetailedText("The details are as follows:")
        msg.setStandardButtons(QMessageBox.Ok)
        msg.exec_()

    def plot_chart(self, chart):
        self.widget = chart
        # chart.plotItem.items[0]
        # self.horizontalLayout.widget = chart
        if self.horizontalLayout.count() > 0:
            self.horizontalLayout.takeAt(0)
        self.horizontalLayout.addWidget(chart)


class MainApp(QtWidgets.QDialog, ui.Ui_Dialog):
    # files = ['s67.s2p', 's68.s2p', 's69.s2p', 's78.s2p']
    files = ['s67.s2p',]


    def __init__(self):
        super().__init__()
        self.setupUi(self)


        self.pushButton_S11.clicked.connect(self.calc_s11_data)
        self.pushButton_S21.clicked.connect(self.calc_s21_data)
        self.pushButton_S12.clicked.connect(self.calc_s12_data)
        self.pushButton_S22.clicked.connect(self.calc_s22_data)
        self.pushButton_VSWR1.clicked.connect(self.calc_VSWR1_data)
        self.pushButton_VSWR2.clicked.connect(self.calc_VSWR2_data)
        self.pushButton_Rx1.clicked.connect(self.calc_Rx1_data)
        self.pushButton_Rx2.clicked.connect(self.calc_Rx1_data)

        self.dialog = Chart()
        self.freq_values = []
        self.y_values = []
        # при нажатии кнопки

    def plot_chart(self):
        chart = pg.PlotWidget(background=pg.mkColor('w'))

        chart.getPlotItem().axes['left']['item'].setPen(pg.mkPen('k', width=2, style=QtCore.Qt.SolidLine))
        chart.getPlotItem().axes['left']['item'].setGrid(255)
        chart.getPlotItem().axes['bottom']['item'].setPen(pg.mkPen('k', width=2, style=QtCore.Qt.SolidLine))
        chart.getPlotItem().axes['bottom']['item'].setGrid(255)

        # symbol = in ['o', 's', 't', 't1', 't2', 't3','d', '+', 'x', 'p', 'h', 'star']
        for i, y_value in enumerate(self.y_values):
            chart.plot(x=self.freq_values, y=y_value, pen=pg.mkPen(MainApp.get_color_by_int(i), width=2))


        # chart_form.show()
        self.dialog.plot_chart(chart)
        self.dialog.show()
        i = 0;

    # -------------------------------------- CLASS METHODS ----------------------#
    @classmethod
    def read_data_file(cls, filename):
        data = {}
        with open(filename) as file:
            content = file.readlines()

        content = content[5:]

        for line in content:
            line = line.split()
            if len(line) == 9:  # case for s2p
                f, *values = line
                data[float(f) / 10 ** 6] = list(map(float, values))

        data = dict(collections.OrderedDict(sorted(data.items())).items())

        freq_values = list(data.keys())
        freq_values = np.asarray(freq_values).astype(np.float)
        values_correct_format = np.asarray(list(data.values())).astype(np.float)

        return freq_values, values_correct_format


    @classmethod
    def get_data_values(cls, column_number, data_files):
        freq_values_out = []
        values = []
        for data_file in data_files:
            freq_values,  values_correct_format = MainApp.read_data_file(data_file)
            if not type(freq_values_out) is np.ndarray:
                freq_values_out = freq_values
            values.append(values_correct_format[:, column_number])
        return freq_values_out, values

    @classmethod
    def get_color_by_int(cls, color_int):
        if color_int == 0:
            return 'k'
        elif color_int == 1:
            return 'r'
        elif color_int == 2:
            return 'b'
        elif color_int == 3:
            return 'g'
        else:
            return 'y'

    @classmethod
    def calc_VSWR(cls, port_number, data_files):

        if port_number == 2:
            freqs, values =  MainApp.get_data_values(6, data_files)  # abs(S22)
        else:
            freqs, values = MainApp.get_data_values(0, data_files)  # abs(S11)

        vswr = []
        for value in values:
            vswr.append((1 + value) / (1 - value))

        return freqs, vswr

    @classmethod
    def calc_Rx(cls, port_number, data_files, z0):

        if port_number == 2:
            freqs, values_abs =  MainApp.get_data_values(6, data_files)  # abs(S22)
            _, values_ang = MainApp.get_data_values(7, data_files)  # abs(S22)
        else:
            freqs, values_abs =  MainApp.get_data_values(0, data_files)  # abs(S22)
            _, values_ang = MainApp.get_data_values(1, data_files)  # abs(S22)

        zx = []
        nprect = np.vectorize(rect)
        for series in range(len(values_abs)):
            value_abs = values_abs[series]
            value_ang_rad = values_ang[series] * np.pi / 180
            complex_ko = nprect(value_abs, value_ang_rad)
            zx.append(z0 * ((1 + complex_ko) / (1 - complex_ko)))

        return freqs, zx

#-------------------------------------- EVENT HANDLERS ----------------------#
    def calc_Rx1_data(self):
        self.freq_values, zx =  MainApp.calc_Rx(1, self.files, 50)
        self.y_values = []
        self.y_values.append(zx[0].real)
        self.y_values.append(zx[0].imag)

        self.plot_chart()

    def calc_Rx2_data(self):
        self.freq_values, self.y_values =  MainApp.calc_Rx(2, self.files, 50)
        self.plot_chart()

    def calc_VSWR1_data(self):
        self.freq_values, self.y_values =  MainApp.calc_VSWR(1, self.files)
        self.plot_chart()

    def calc_VSWR2_data(self):
        self.freq_values, self.y_values = MainApp.calc_VSWR(2, self.files)
        self.plot_chart()

    def calc_s11_data(self):
        self.y_values = []
        self.freq_values, self.y_values = MainApp.get_data_values(0, self.files)
        self.plot_chart()

    def calc_s12_data(self):
        self.y_values = []
        self.freq_values, self.y_values = MainApp.get_data_values(2, self.files)
        self.plot_chart()

    def calc_s21_data(self):
        self.y_values = []
        self.freq_values, self.y_values = MainApp.get_data_values(4, self.files)
        self.plot_chart()

    def calc_s22_data(self):
        self.y_values = []
        self.freq_values, self.y_values = MainApp.get_data_values(6, self.files)
        self.plot_chart()

    # def browse_folder(self):
    #     files = QtWidgets.QFileDialog.getOpenFileName(self, 'Выберите файл')
    #     #
    #     # for f in files:
    #     #     self.listWidget.addItem(f)
    #
    #
    #
    #     for key, value in ordered_data.items():
    #         x_dates.append(key)
    #         y_dates.append(value[0])
    #
    #     x_dates = np.asarray(x_dates).astype(np.float)
    #     y_dates = np.asarray(y_dates).astype(np.float)
    #     axes = pg.AxisItem(orientation='right')
    #     axes.setPen(color=pg.mkColor('b'))


        # self.horizontalLayout


def main():
    app = QtWidgets.QApplication(sys.argv)
    window = MainApp()
    window.show()
    app.exec_()



if __name__ == '__main__':
    main()
