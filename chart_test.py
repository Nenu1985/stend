import os
import sys
import collections
from docx.opc.exceptions import PackageNotFoundError
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import QMessageBox
import pyqtgraph as pg
import PGQ_image_exporter # modified pyqtgraph Exporter!
import numpy as np
import ui
import form
from cmath import rect
from docx import Document
import datetime
import cv2


class Chart(QtWidgets.QDialog, form.Ui_Dialog):

    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.files_names = []

        self.openButton.clicked.connect(self.plot_to_word)

        self.spinBox_x_max.valueChanged.connect(self.spinBox_x_max_valueChanged)
        self.spinBox_x_min.valueChanged.connect(self.spinBox_x_max_valueChanged)
        self.doubleSpinBox_y_max.valueChanged.connect(self.doubleSpinBox_y_max_valueChanged)
        self.doubleSpinBox_y_min.valueChanged.connect(self.doubleSpinBox_y_max_valueChanged)

    def plot_to_word(self):
        FILE_EXIST_FILE = False
        # create an exporter instance, as an argument give it
        # the item you wish to export
        exporter = PGQ_image_exporter.PQG_ImageExporter(self.widget.plotItem)

        # set export parameters if needed
        # exporter.parameters()['width'] = 600

        # save to file
        exporter.export('fileName.jpg')

        word_file_name = self.lineEdit_word_file_path.text()
        try:
            document = Document(word_file_name)
            FILE_EXIST_FILE = True
        except PackageNotFoundError as e:
            try:
                document = Document()
            except Exception as e:
                print('Error while creating word template: {}'.format(e))
        except Exception as e:
            print('Error while saving to word {}'.format(e))
            return

        document.add_paragraph(
            'График добавлен. {}'.format(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        document.add_picture('fileName.jpg')

        for file_name in self.files_names:
            document.add_paragraph(file_name)
        document.add_page_break()


        # document.save(word_file_name)
        try:
            document.save(word_file_name)
        except PackageNotFoundError as e:
            print('Package not found exception while saving to word')
            return
        except Exception as e:
            print('Error while saving to word {}'.format(e))
            return

        Chart.display_ok_message('Word saved!')

    @classmethod
    def display_ok_message(cls, message):
        msg = QtWidgets.QMessageBox()
        msg.setIcon(QMessageBox.Information)

        msg.setText(message)
        msg.setInformativeText("График добавлен в Word!")
        msg.setWindowTitle("")
        # msg.setDetailedText("The details are as follows:")
        msg.setStandardButtons(QMessageBox.Ok)
        msg.exec_()

    def plot_chart(self, chart):
        # self.widget = chart

        # chart.plotItem.items[0]
        self.horizontalLayout.widget = chart
        while self.horizontalLayout.count() > 0:
            self.horizontalLayout.takeAt(0)
        self.horizontalLayout.addWidget(chart)
        #
        # graph = next((x for x in self.widget.items() if x is pg.PlotItem), None)
        # x_range = [int(i) for i in self.widget.plotItem.viewRange()[0]]
        # y_range = self.widget.plotItem.viewRange()[1]
        #
        # self.grid_range_settings(x_range, 'bottom')
        #
        # self.spinBox_x_min.setValue(x_range[0])
        # self.spinBox_x_max.setValue(x_range[1])
        # self.doubleSpinBox_y_min.setValue(y_range[0])
        # self.doubleSpinBox_y_max.setValue(y_range[1])

    def grid_range_settings(self, values_range, axis):
        start = values_range[0]
        end = values_range[1]
        dx = int((end-start)/10)
        delta = [(value, str(value)) for value in list(range(start, end+dx, dx))]
        ax = self.widget.getAxis(axis)
        ax.setTicks([delta, []])


    def spinBox_x_max_valueChanged(self):
        self.grid_range_settings(
            [self.spinBox_x_min.value(), self.spinBox_x_max.value()],
            'bottom'
        )
        self.widget.plotItem.setRange(
            xRange=[self.spinBox_x_min.value(), self.spinBox_x_max.value()],
            padding=0,
        )
    def doubleSpinBox_y_max_valueChanged(self):
        self.widget.plotItem.setRange(
            yRange=[self.doubleSpinBox_y_min.value(), self.doubleSpinBox_y_max.value()],
            padding=0,
        )