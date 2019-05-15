import os
import sys
import collections
from docx.opc.exceptions import PackageNotFoundError
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import QMessageBox
import pyqtgraph as pg
import PGQ_image_exporter # modified pyqtgraph Exporter!
import numpy as np
import ui
import form
from cmath import rect
from docx import Document
import datetime
import cv2
from PyQt5.QtGui import QFont


class CustomViewBox(pg.ViewBox):
    # cross hair



    def __init__(self, parent_chart_dialog, *args, **kwds):
        pg.ViewBox.__init__(self, *args, **kwds)
        self.setMouseMode(self.RectMode)
        self.chart_dialog_form = parent_chart_dialog
        self.label = pg.LabelItem(text='Hello!', justify='left')
        # self.addItem(self.label)
        self.proxy = ''
        self.vLine = pg.InfiniteLine(angle=90, movable=False)
        self.hLine = pg.InfiniteLine(angle=0, movable=False)
        # self.addItem(self.vLine)
        # self.addItem(self.hLine)

    ## reimplement right-click to zoom out
    def mouseClickEvent(self, ev):
        if ev.button() == QtCore.Qt.RightButton:
            self.autoRange()
            self.chart_dialog_form.spinBox_x_max_valueChanged()
            self.chart_dialog_form.doubleSpinBox_y_max_valueChanged()
            self.chart_dialog_form.plot_chart()

    def subscribe_to_mouse_event(self):
        self.proxy = pg.SignalProxy(self.chart_dialog_form.chart.scene().sigMouseMoved, rateLimit=60, slot=self.mouseMoved)

    def mouseDragEvent(self, ev):
        if ev.button() == QtCore.Qt.RightButton:
            ev.ignore()
        else:
            pg.ViewBox.mouseDragEvent(self, ev)

    def wheelEvent(self, ev, axis=None):
        ev.ignore()

    def mouseMoved(self, evt):
        pos = evt[0]  ## using signal proxy turns original arguments into a tuple
        if self.chart_dialog_form.chart.sceneBoundingRect().contains(pos):
            mousePoint = self.mapSceneToView(pos)
            index = int(mousePoint.x())
            if index > 0 and index < 1000:
                self.label.setText(
                    "<span style='font-size: 12pt'>"
                        "<span style='color: blue'>x=%0.1f</span>, <br>  "
                        "<span style='color: red' >y=%0.3f</span>, <br>"
                        # "<span style='color: green'>y2=%0.1f"
                    "</span>" % (
                    mousePoint.x(), mousePoint.y()))
            self.vLine.setPos(mousePoint.x())
            self.hLine.setPos(mousePoint.y())


class Chart(QtWidgets.QDialog, form.Ui_Dialog):

    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.chart = pg.PlotItem()
        self.view = pg.GraphicsView()
        self.gr_layout = pg.GraphicsLayout()
        self.files_names = []

        # self.plot_chart_init()

        self.openButton.clicked.connect(self.plot_to_word)

        self.spinBox_x_max.valueChanged.connect(self.spinBox_x_max_valueChanged)
        self.spinBox_x_min.valueChanged.connect(self.spinBox_x_max_valueChanged)
        self.doubleSpinBox_y_max.valueChanged.connect(self.doubleSpinBox_y_max_valueChanged)
        self.doubleSpinBox_y_min.valueChanged.connect(self.doubleSpinBox_y_max_valueChanged)

    def plot_chart_init(self):
        vb = CustomViewBox(self)

        self.gr_layout = pg.GraphicsLayout()
        # set left margin
        self.gr_layout.setContentsMargins(50., 0., 0., 0.)
        # layout.addViewBox(0,0)
        self.view = pg.GraphicsView(background=pg.mkColor('w'))
        self.view.setCentralItem(self.gr_layout)
        #
        self.chart = pg.PlotItem(enableMenu=False, viewBox=vb)
        self.gr_layout.addItem(self.chart, 0, 0)
        # self.chart = self.gr_layout.addPlot(0, 0, enableMenu=False, viewBox=vb)
        self.gr_layout.addItem(vb.label, 0, 1)
        self.chart.addItem(vb.vLine)
        self.chart.addItem(vb.hLine)
        # legend = layout.addLabel('_________________',0,1)
        vb.subscribe_to_mouse_event()

        bottom_axes = self.chart.getAxis('bottom')
        left_axes = self.chart.getAxis('left')
        right_axes = self.chart.getAxis('right')
        top_axes = self.chart.getAxis('top')

        self.gr_layout.addItem(self.chart.addLegend(size=(200, 100), offset=(1380, 10)), 0, 1)

        left_axes.setPen(pg.mkPen('k', width=2, style=QtCore.Qt.SolidLine))
        left_axes.setLabel('  ')

        self.chart.getAxis('left').setGrid(150)
        bottom_axes.setPen(pg.mkPen('k', width=2, style=QtCore.Qt.SolidLine))
        bottom_axes.setGrid(150)
        bottom_axes.setLabel('Частота', units='МГц')

        top_axes.show()
        top_axes.setPen(pg.mkPen('k', width=2))
        top_axes.setTicks([])

        right_axes.setLabel('   ', color='#fff')
        right_axes.setPen(pg.mkPen('k', width=2))
        right_axes.setTicks([])
        right_axes.show()
        # chart_pw.showAxis('right')
        # results_graph(PlotWidget).setLabels(left="Frequency", bottom="Scores")
        bottom_axes.labelStyle = {'font-size': '20pt', 'color': '#000'}

        tick_font = QFont()
        tick_font.setPixelSize(20)
        axis_style = {
            'tickFont': tick_font,
            'tickLength': 10,
            'tickTextOffset': 15,
            'tickTextHeight': 20,
            'tickTextWidth': 20,
            # 'stopAxisAtTick': (False, False),
        }

        bottom_axes.setTickFont(tick_font)
        left_axes.setTickFont(tick_font)

        bottom_axes.setStyle(**axis_style)
        left_axes.setStyle(**axis_style)

        # chart_pw.setBorder('k', width=2)

        # self.chart.plot(border=pg.mkPen('k', width=5), symbol='o')
        # self.chart.plot(x=[1,2,3,4,5], y = [10,20,30,40,50], border=pg.mkPen('k', width=5), symbol='o')
        while self.horizontalLayout.count() > 0:
            self.horizontalLayout.takeAt(0)
        self.horizontalLayout.addWidget(self.view)
        # --self.chart.setTitle(self.title, **{'color': '#000', 'size': '14pt'})

    def get_chart_plot_items_number(self):
        number = 0
        for plot_item in self.chart.items:
            if type(plot_item) is pg.PlotCurveItem:
                number += 1
        return number

    def delete_plot_items(self):
        self.plot_chart_init()

    def plot_to_word(self):
        FILE_EXIST_FILE = False
        # create an exporter instance, as an argument give it
        # the item you wish to export
        exporter = PGQ_image_exporter.PQG_ImageExporter(self.chart)

        # set export parameters if needed
        # exporter.parameters()['width'] = 600

        # save to file
        exporter.export('fileName.jpg')

        word_file_name = self.lineEdit_word_file_path.text()
        try:
            document = Document(word_file_name)
            FILE_EXIST_FILE = True
        except PackageNotFoundError as e:
            try:
                document = Document()
            except Exception as e:
                print('Error while creating word template: {}'.format(e))
        except Exception as e:
            print('Error while saving to word {}'.format(e))
            return

        document.add_paragraph(
            'График добавлен. {}'.format(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        document.add_picture('fileName.jpg')

        for file_name in self.files_names:
            document.add_paragraph(file_name)
        document.add_page_break()

        try:
            document.save(word_file_name)
        except PackageNotFoundError as e:
            print('Package not found exception while saving to word')
            return
        except Exception as e:
            print('Error while saving to word {}'.format(e))
            return

        Chart.display_ok_message('Word saved!')

    @classmethod
    def display_ok_message(cls, message):
        msg = QtWidgets.QMessageBox()
        msg.setIcon(QMessageBox.Information)

        msg.setText(message)
        msg.setInformativeText("График добавлен в Word!")
        msg.setWindowTitle("")
        # msg.setDetailedText("The details are as follows:")
        msg.setStandardButtons(QMessageBox.Ok)
        msg.exec_()

    def plot_chart(self):

        x_min = 9999;
        x_max = -9999;
        y_max = -9999;
        y_min = 9999;
        for plot_item in self.chart.items:
            if type(plot_item) is pg.PlotCurveItem or type(plot_item) is pg.PlotDataItem:
                if min(plot_item.xData) < x_min:
                    x_min = min(plot_item.xData)
                if max(plot_item.xData) > x_max:
                    x_max = max(plot_item.xData)
                if min(plot_item.yData) < y_min:
                    y_min = min(plot_item.yData)
                if max(plot_item.yData) > y_max:
                    y_max = max(plot_item.yData)

        self.spinBox_x_min.setValue(x_min)
        self.spinBox_x_max.setValue(x_max)
        self.doubleSpinBox_y_min.setValue(y_min)
        self.doubleSpinBox_y_max.setValue(y_max)
        self.spinBox_x_max_valueChanged()
        self.doubleSpinBox_y_max_valueChanged()

    def grid_range_settings_x(self, values_range):
        start = values_range[0]
        end = values_range[1]

        dx = int((end-start)/8)
        if dx == 0:
            print('Слишком маленький диапазон')
            return
        delta = [(value, str(value)) for value in list(range(start, end-dx, dx))]
        delta.append((end, str(end)))
        ax = self.chart.getAxis('bottom')
        ax.setTicks([delta, []])

    def grid_range_settings_y(self, values_range):
        start = values_range[0]
        end = values_range[1]
        dy = (end-start)/6
        delta = [(value, '{:.2f}'.format(value)) for value in list(drange(start, end, dy))]
        delta.append((end, '{:.2f}'.format(end)))
        ay = self.chart.getAxis('left')
        ay.setTicks([delta, []])

    def spinBox_x_max_valueChanged(self):
        self.grid_range_settings_x(
            [self.spinBox_x_min.value(), self.spinBox_x_max.value()],
        )
        self.chart.setRange(
            xRange=[self.spinBox_x_min.value(), self.spinBox_x_max.value()],
            padding=0,
        )
    def doubleSpinBox_y_max_valueChanged(self):
        self.grid_range_settings_y(
            [self.doubleSpinBox_y_min.value(), self.doubleSpinBox_y_max.value()],
        )
        self.chart.setRange(
            yRange=[self.doubleSpinBox_y_min.value(), self.doubleSpinBox_y_max.value()],
            padding=0,
        )

def drange(start, stop, step):
        r = start
        while r < stop:
            yield r
            r += step
